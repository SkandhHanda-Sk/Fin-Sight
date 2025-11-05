#!/usr/bin/env python3
"""
Multi-Format Text Extractor (Backend Module)
Extracts text from XLSX, XLS, PDF, and DOCX files with a focus on preserving 
table structure and OCR support for scanned PDFs.
"""

import os
import sys
import logging
import magic
from typing import Dict, Callable, Set
import pandas as pd
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image

# Configure logging for clear, standardized error and info messages
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# --- Configuration Constants ---
TABLE_SEPARATOR = "\t"
OCR_RESOLUTION = 300
MIME_TYPE_MAP = {
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
    'application/vnd.ms-excel': '.xls',
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
}

class ExtractionError(Exception):
    """Custom exception for user-facing extraction failures."""
    pass

class TextExtractor:
    """
    Unified text extractor for various document formats.
    Designed for backend processing with security and robustness in mind.
    """
    def __init__(self):
        self.extractors: Dict[str, Callable[[str], str]] = {
            '.xlsx': self._extract_from_excel,
            '.xls': self._extract_from_excel,
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
        }
        self.supported_mime_types: Set[str] = set(MIME_TYPE_MAP.keys())

    def _get_file_format_from_mime(self, file_path: str) -> str:
        """Detects the file format using its MIME type for security."""
        try:
            mime_type = magic.from_file(file_path, mime=True)
            if mime_type in self.supported_mime_types:
                return MIME_TYPE_MAP[mime_type]
            return ""
        except Exception as e:
            logging.error(f"Could not determine MIME type for {file_path}: {e}")
            return ""

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files, preserving sheet and row structure."""
        try:
            text_parts = []
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                if df.empty:
                    continue
                text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
                for _, row in df.iterrows():
                    row_text = [str(cell).strip().replace('\n', ' ') for cell in row if pd.notna(cell)]
                    if row_text:
                        text_parts.append(TABLE_SEPARATOR.join(row_text))
            return "\n".join(text_parts)
        except Exception as e:
            raise ExtractionError(f"Failed to process Excel file {os.path.basename(file_path)}.") from e

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text and tables from PDF files, with an OCR fallback."""
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text_content.append(f"\n--- Page {i}/{len(pdf.pages)} ---\n")
                    page_text = page.extract_text() or ""
                    
                    if len(page_text.strip()) < 20:
                        logging.info(f"Page {i} of {os.path.basename(file_path)} has minimal text. Attempting OCR.")
                        try:
                            img = page.to_image(resolution=OCR_RESOLUTION)
                            ocr_text = pytesseract.image_to_string(img.original)
                            if ocr_text.strip():
                                text_content.append("\n--- OCR Extracted Text (Scanned Page) ---\n")
                                text_content.append(ocr_text)
                            else:
                                text_content.append("[Warning: Page appears to be blank or an image with no text found by OCR.]")
                        except Exception as ocr_error:
                            logging.error(f"OCR failed on page {i} of {file_path}: {ocr_error}")
                            text_content.append("[Error: OCR processing failed for this page.]")
                    else:
                        text_content.append(page_text)

                    tables = page.extract_tables()
                    if tables:
                        text_content.append("\n--- Tables on Page ---")
                        for table in tables:
                            text_content.append("\n-- Table Start --\n")
                            for row in table:
                                clean_row = [str(cell).strip().replace('\n', ' ') if cell is not None else "" for cell in row]
                                text_content.append(TABLE_SEPARATOR.join(clean_row))
                            text_content.append("\n-- Table End --\n")
            return "\n".join(text_content)
        except pdfplumber.errors.PasswordRequired:
            raise ExtractionError("PDF file is password-protected.")
        except Exception as e:
            raise ExtractionError(f"Failed to process PDF file {os.path.basename(file_path)}.") from e

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text and tables from DOCX files."""
        try:
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            if doc.tables:
                text_parts.append("\n--- Tables ---")
                for table in doc.tables:
                    text_parts.append("\n-- Table Start --\n")
                    for row in table.rows:
                        row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        text_parts.append(TABLE_SEPARATOR.join(row_text))
                    text_parts.append("\n-- Table End --\n")
            return "\n".join(text_parts)
        except Exception as e:
            raise ExtractionError(f"Failed to process DOCX file {os.path.basename(file_path)}.") from e

    def extract_text(self, file_path: str) -> str:
        """Main public method. Validates and extracts text from a supported file."""
        if not os.path.exists(file_path):
            return f"[Error: File not found at path: {file_path}]"

        file_format = self._get_file_format_from_mime(file_path)
        if not file_format:
            logging.warning(f"Could not verify MIME type for {os.path.basename(file_path)}. Falling back to extension.")
            _, ext = os.path.splitext(file_path)
            file_format = ext.lower()

        extractor_func = self.extractors.get(file_format)
        if not extractor_func:
            return f"[Error: Unsupported file format. Please upload a XLSX, XLS, PDF, or DOCX file.]"

        logging.info(f"Extracting text from '{os.path.basename(file_path)}' using {file_format} extractor...")
        try:
            return extractor_func(file_path)
        except ExtractionError as e:
            logging.error(f"Extraction failed for {os.path.basename(file_path)}: {e}")
            return f"[Error: {e}]"
        except Exception as e:
            logging.critical(f"An unexpected error occurred during extraction of {os.path.basename(file_path)}: {e}", exc_info=True)
            return "[Error: An unexpected server error occurred. Please contact support.]"

    @staticmethod
    def save_text(text: str, output_file: str):
        """Saves the given text to a file."""
        try:
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)
        except IOError as e:
            logging.error(f"Failed to save text to {output_file}: {e}")
            raise

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python text_extractor.py <path_to_input_file> [path_to_output_file]")
        sys.exit(1)

    input_file = sys.argv[1]
    
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = f"output/{base_name}_extracted.txt"

    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    print(f"Processing {input_file}...")
    extractor = TextExtractor()
    extracted_text = extractor.extract_text(input_file)

    if not extracted_text.startswith("[Error:"):
        try:
            extractor.save_text(extracted_text, output_file)
            print(f"✅ Successfully extracted text to {output_file}")
        except Exception as e:
            print(f"❌ Error saving file: {e}")
    else:
        print(f"❌ Extraction failed. Reason: {extracted_text}")